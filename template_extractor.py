"""
Template Extractor Module

Extracts formatting metadata from PDF and DOCX templates including:
- Headers and footers
- Page breaks
- Fonts and text sizes
- Numbering and bullets
- Tables
- Styles and formatting
"""

import os
from typing import Dict, List, Any, Optional
from pathlib import Path

try:
    from pdfplumber import PDF
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False


class TemplateExtractor:
    """Extract formatting metadata from PDF and DOCX templates"""

    def __init__(self):
        self.supported_formats = []
        if PDFPLUMBER_AVAILABLE or PYMUPDF_AVAILABLE:
            self.supported_formats.append("pdf")
        if PYTHON_DOCX_AVAILABLE:
            self.supported_formats.append("docx")

    def extract_template(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Extract template metadata from a file

        Args:
            file_path: Path to the template file
            file_type: Type of file (pdf or docx)

        Returns:
            Dictionary containing extracted metadata
        """
        if file_type.lower() == "pdf":
            return self._extract_pdf(file_path)
        elif file_type.lower() == "docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file"""
        metadata = {
            "headers": [],
            "footers": [],
            "page_breaks": [],
            "fonts": {},
            "text_sizes": {},
            "tables": [],
            "styles": {},
            "numbering": [],
            "bullets": [],
            "page_count": 0,
            "margins": {},
            "page_size": {}
        }

        # Try pdfplumber first
        if PDFPLUMBER_AVAILABLE:
            try:
                with PDF.open(file_path) as pdf:
                    metadata["page_count"] = len(pdf.pages)

                    for page_num, page in enumerate(pdf.pages, 1):
                        # Extract text with positioning
                        words = page.extract_words()

                        # Identify headers (typically at top of page)
                        page_height = page.height
                        header_threshold = page_height * 0.1  # Top 10% of page

                        for word in words:
                            if word["top"] < header_threshold:
                                metadata["headers"].append({
                                    "text": word["text"],
                                    "font": word.get("fontname", "Unknown"),
                                    "size": word.get("size", 12),
                                    "page": page_num,
                                    "position": {"top": word["top"], "left": word["left"]}
                                })

                        # Extract tables
                        tables = page.extract_tables()
                        for table_idx, table in enumerate(tables):
                            table_metadata = {
                                "page": page_num,
                                "rows": len(table),
                                "columns": len(table[0]) if table else 0,
                                "cells": []
                            }

                            # Extract cell formatting
                            bbox = page.find_tables()[table_idx].bbox if page.find_tables() else None
                            if bbox:
                                table_metadata["bbox"] = {
                                    "x0": bbox[0],
                                    "y0": bbox[1],
                                    "x1": bbox[2],
                                    "y1": bbox[3]
                                }

                            metadata["tables"].append(table_metadata)

                        # Collect font information
                        for word in words:
                            font_name = word.get("fontname", "Unknown")
                            font_size = word.get("size", 12)

                            if font_name not in metadata["fonts"]:
                                metadata["fonts"][font_name] = []
                            metadata["fonts"][font_name].append(font_size)

                            if font_size not in metadata["text_sizes"]:
                                metadata["text_sizes"][font_size] = []
                            metadata["text_sizes"][font_size].append(font_name)

                        # Page break (end of page)
                        if page_num < len(pdf.pages):
                            metadata["page_breaks"].append({
                                "page": page_num,
                                "type": "page_break"
                            })

                    # Get page size from first page
                    if pdf.pages:
                        first_page = pdf.pages[0]
                        metadata["page_size"] = {
                            "width": first_page.width,
                            "height": first_page.height
                        }

                return metadata
            except Exception as e:
                # Fall back to PyMuPDF if pdfplumber fails
                if PYMUPDF_AVAILABLE:
                    return self._extract_pdf_pymupdf(file_path)
                raise e

        # Use PyMuPDF as fallback
        if PYMUPDF_AVAILABLE:
            return self._extract_pdf_pymupdf(file_path)

        raise ImportError("No PDF extraction library available. Install pdfplumber or PyMuPDF.")

    def _extract_pdf_pymupdf(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata using PyMuPDF"""
        metadata = {
            "headers": [],
            "footers": [],
            "page_breaks": [],
            "fonts": {},
            "text_sizes": {},
            "tables": [],
            "styles": {},
            "numbering": [],
            "bullets": [],
            "page_count": 0,
            "margins": {},
            "page_size": {}
        }

        doc = fitz.open(file_path)
        metadata["page_count"] = len(doc)

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Get page dimensions
            rect = page.rect
            metadata["page_size"] = {
                "width": rect.width,
                "height": rect.height
            }

            # Extract text blocks with formatting
            blocks = page.get_text("dict")

            header_threshold = rect.height * 0.1

            for block in blocks.get("blocks", []):
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            font_name = span.get("font", "Unknown")
                            font_size = span.get("size", 12)
                            text = span.get("text", "")
                            bbox = span.get("bbox", [0, 0, 0, 0])

                            # Check if header
                            if bbox[1] < header_threshold:
                                metadata["headers"].append({
                                    "text": text,
                                    "font": font_name,
                                    "size": font_size,
                                    "page": page_num + 1,
                                    "position": {"top": bbox[1], "left": bbox[0]}
                                })

                            # Collect font info
                            if font_name not in metadata["fonts"]:
                                metadata["fonts"][font_name] = []
                            metadata["fonts"][font_name].append(font_size)

                            if font_size not in metadata["text_sizes"]:
                                metadata["text_sizes"][font_size] = []
                            metadata["text_sizes"][font_size].append(font_name)

            # Extract tables
            tables = page.find_tables()
            for table in tables:
                table_metadata = {
                    "page": page_num + 1,
                    "rows": table.row_count,
                    "columns": table.col_count,
                    "bbox": {
                        "x0": table.bbox.x0,
                        "y0": table.bbox.y0,
                        "x1": table.bbox.x1,
                        "y1": table.bbox.y1
                    }
                }
                metadata["tables"].append(table_metadata)

            # Page break
            if page_num < len(doc) - 1:
                metadata["page_breaks"].append({
                    "page": page_num + 1,
                    "type": "page_break"
                })

        doc.close()
        return metadata

    def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file"""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX extraction")

        doc = Document(file_path)
        metadata = {
            "headers": [],
            "footers": [],
            "page_breaks": [],
            "fonts": {},
            "text_sizes": {},
            "tables": [],
            "styles": {},
            "numbering": [],
            "bullets": [],
            "paragraphs": [],
            "sections": []
        }

        # Extract paragraph formatting
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_metadata = {
                "index": para_idx,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else None,
                "alignment": str(paragraph.alignment) if paragraph.alignment else None,
                "runs": []
            }

            # Extract run-level formatting
            for run in paragraph.runs:
                run_metadata = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name if run.font.name else None,
                    "font_size": run.font.size.pt if run.font.size else None
                }

                para_metadata["runs"].append(run_metadata)

                # Collect font information
                if run_metadata["font_name"]:
                    if run_metadata["font_name"] not in metadata["fonts"]:
                        metadata["fonts"][run_metadata["font_name"]] = []
                    if run_metadata["font_size"]:
                        metadata["fonts"][run_metadata["font_name"]].append(run_metadata["font_size"])

                if run_metadata["font_size"]:
                    if run_metadata["font_size"] not in metadata["text_sizes"]:
                        metadata["text_sizes"][run_metadata["font_size"]] = []
                    if run_metadata["font_name"]:
                        metadata["text_sizes"][run_metadata["font_size"]].append(run_metadata["font_name"])

            # Check for headers (styles starting with "Heading")
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                metadata["headers"].append({
                    "text": paragraph.text,
                    "style": paragraph.style.name,
                    "level": int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1,
                    "index": para_idx
                })

            # Check for numbering/bullets
            if paragraph.style:
                if "List" in paragraph.style.name or "Bullet" in paragraph.style.name:
                    metadata["bullets"].append({
                        "text": paragraph.text,
                        "style": paragraph.style.name,
                        "index": para_idx
                    })

            metadata["paragraphs"].append(para_metadata)

        # Extract table formatting
        for table_idx, table in enumerate(doc.tables):
            table_metadata = {
                "index": table_idx,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": []
            }

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_metadata = {
                        "row": row_idx,
                        "column": col_idx,
                        "text": cell.text,
                        "paragraphs": []
                    }

                    for para in cell.paragraphs:
                        para_info = {
                            "text": para.text,
                            "style": para.style.name if para.style else None,
                            "alignment": str(para.alignment) if para.alignment else None
                        }
                        cell_metadata["paragraphs"].append(para_info)

                    table_metadata["cells"].append(cell_metadata)

            metadata["tables"].append(table_metadata)

        # Extract section information
        for section in doc.sections:
            section_metadata = {
                "page_width": section.page_width.pt if section.page_width else None,
                "page_height": section.page_height.pt if section.page_height else None,
                "left_margin": section.left_margin.pt if section.left_margin else None,
                "right_margin": section.right_margin.pt if section.right_margin else None,
                "top_margin": section.top_margin.pt if section.top_margin else None,
                "bottom_margin": section.bottom_margin.pt if section.bottom_margin else None
            }
            metadata["sections"].append(section_metadata)

        # Extract style information
        for style in doc.styles:
            style_metadata = {
                "name": style.name,
                "type": str(style.type),
                "font_name": style.font.name if style.font.name else None,
                "font_size": style.font.size.pt if style.font.size else None,
                "bold": style.font.bold,
                "italic": style.font.italic
            }
            metadata["styles"][style.name] = style_metadata

        return metadata
