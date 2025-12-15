"""
PDF Generator Module

Generates PDFs using template metadata to match the original formatting exactly.
"""

import os
from typing import Dict, List, Any, Optional
from pathlib import Path
import tempfile

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.units import inch, pt
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, KeepTogether
    from reportlab.platypus.flowables import Image
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class PDFGenerator:
    """Generate PDFs matching template formatting"""

    def __init__(self, temp_dir: Optional[str] = None):
        """
        Initialize PDF generator

        Args:
            temp_dir: Temporary directory for generated PDFs (defaults to system temp)
        """
        # Use environment variable or system temp for Docker compatibility
        if temp_dir is None:
            temp_dir = os.getenv("PDF_TEMP_DIR", tempfile.gettempdir())

        # Create subdirectory for our PDFs
        self.temp_dir = Path(temp_dir) / "pdf_template_formatter"
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self._registered_fonts = {}

    def generate_pdf(
        self,
        content: str,
        template_metadata: Dict[str, Any],
        template_file_path: Optional[str] = None,
        output_name: str = "output.pdf"
    ) -> str:
        """
        Generate a PDF from content using template metadata

        Args:
            content: The content to format
            template_metadata: Extracted template metadata
            template_file_path: Path to original template (for reference)
            output_name: Name for output PDF

        Returns:
            Path to generated PDF file
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF generation")

        # Determine page size
        page_size = template_metadata.get("page_size", {})
        if page_size:
            page_width = page_size.get("width", 8.5 * inch)
            page_height = page_size.get("height", 11 * inch)
        else:
            page_width, page_height = letter

        # Create PDF document
        output_path = os.path.join(self.temp_dir, output_name)
        doc = SimpleDocTemplate(
            output_path,
            pagesize=(page_width, page_height),
            rightMargin=template_metadata.get("margins", {}).get("right", 1 * inch) or 1 * inch,
            leftMargin=template_metadata.get("margins", {}).get("left", 1 * inch) or 1 * inch,
            topMargin=template_metadata.get("margins", {}).get("top", 1 * inch) or 1 * inch,
            bottomMargin=template_metadata.get("margins", {}).get("bottom", 1 * inch) or 1 * inch
        )

        # Build story (content elements)
        story = []

        # Register fonts from template
        fonts = template_metadata.get("fonts", {})
        self._register_template_fonts(fonts)

        # Parse content and apply formatting
        content_elements = self._parse_content_with_formatting(
            content,
            template_metadata
        )

        story.extend(content_elements)

        # Build PDF
        doc.build(story)

        return output_path

    def _register_template_fonts(self, fonts: Dict[str, List[float]]):
        """Register fonts found in template"""
        # Try to register common fonts
        common_fonts = {
            "Times-Roman": "Times-Roman",
            "Times-Bold": "Times-Bold",
            "Helvetica": "Helvetica",
            "Helvetica-Bold": "Helvetica-Bold",
            "Courier": "Courier",
            "Courier-Bold": "Courier-Bold"
        }

        for font_name in fonts.keys():
            # Normalize font name
            normalized = font_name.split(",")[0].strip() if "," in font_name else font_name

            # Map to ReportLab font names
            if normalized in common_fonts:
                self._registered_fonts[font_name] = common_fonts[normalized]
            else:
                # Default to Helvetica
                self._registered_fonts[font_name] = "Helvetica"

    def _parse_content_with_formatting(
        self,
        content: str,
        template_metadata: Dict[str, Any]
    ) -> List:
        """Parse content and create formatted elements matching template"""
        elements = []

        # Get default styles
        styles = getSampleStyleSheet()

        # Create custom styles based on template
        template_styles = template_metadata.get("styles", {})
        custom_styles = {}

        # Map template styles to ReportLab styles
        for style_name, style_info in template_styles.items():
            font_name = self._registered_fonts.get(
                style_info.get("font_name", "Helvetica"),
                "Helvetica"
            )
            font_size = style_info.get("font_size", 12)

            custom_styles[style_name] = ParagraphStyle(
                name=style_name,
                parent=styles["Normal"],
                fontName=font_name,
                fontSize=font_size,
                spaceBefore=6,
                spaceAfter=6
            )

        # Get header styles from template
        headers = template_metadata.get("headers", [])
        header_styles = {}
        if headers:
            # Group headers by level
            for header in headers:
                level = header.get("level", 1)
                font_name = self._registered_fonts.get(
                    header.get("font", "Helvetica-Bold"),
                    "Helvetica-Bold"
                )
                font_size = header.get("size", 14 + (4 - level) * 2)

                if level not in header_styles:
                    header_styles[level] = ParagraphStyle(
                        name=f"Heading{level}",
                        parent=styles["Heading1"],
                        fontName=font_name,
                        fontSize=font_size,
                        spaceBefore=12,
                        spaceAfter=6,
                        textColor=colors.black
                    )

        # Parse content into paragraphs
        paragraphs = content.split("\n\n")

        for para_text in paragraphs:
            if not para_text.strip():
                elements.append(Spacer(1, 6))
                continue

            # Check if it's a markdown table
            if self._is_markdown_table(para_text):
                table_data = self._parse_markdown_table(para_text)
                if table_data:
                    table = self._create_table_from_data(table_data, template_metadata)
                    elements.append(table)
                    elements.append(Spacer(1, 12))
                continue

            # Check if it's a header (starts with #)
            if para_text.strip().startswith("#"):
                level = len(para_text) - len(para_text.lstrip("#"))
                level = min(level, 6)

                text = para_text.lstrip("#").strip()
                style = header_styles.get(level, styles["Heading1"])
                elements.append(Paragraph(text, style))
                elements.append(Spacer(1, 6))

            # Check if it's a bullet list
            elif para_text.strip().startswith("-") or para_text.strip().startswith("*"):
                lines = para_text.split("\n")
                for line in lines:
                    if line.strip().startswith("-") or line.strip().startswith("*"):
                        text = line.lstrip("-*").strip()
                        bullet_style = ParagraphStyle(
                            name="Bullet",
                            parent=styles["Normal"],
                            leftIndent=20,
                            bulletIndent=10,
                            bulletText="•"
                        )
                        elements.append(Paragraph(f"• {text}", bullet_style))
                elements.append(Spacer(1, 6))

            # Check if it's a numbered list
            elif any(line.strip()[0].isdigit() and "." in line[:5] for line in para_text.split("\n") if line.strip()):
                lines = para_text.split("\n")
                for idx, line in enumerate(lines, 1):
                    if line.strip():
                        # Remove number prefix if present
                        text = line.strip()
                        if text[0].isdigit():
                            parts = text.split(".", 1)
                            if len(parts) > 1:
                                text = parts[1].strip()

                        numbered_style = ParagraphStyle(
                            name="Numbered",
                            parent=styles["Normal"],
                            leftIndent=20
                        )
                        elements.append(Paragraph(f"{idx}. {text}", numbered_style))
                elements.append(Spacer(1, 6))

            # Regular paragraph
            else:
                # Use default font from template
                default_font = list(template_metadata.get("fonts", {}).keys())[0] if template_metadata.get("fonts") else "Helvetica"
                default_size = template_metadata.get("text_sizes", {})
                default_size = list(default_size.keys())[0] if default_size else 12

                para_style = ParagraphStyle(
                    name="Normal",
                    parent=styles["Normal"],
                    fontName=self._registered_fonts.get(default_font, "Helvetica"),
                    fontSize=default_size,
                    spaceBefore=6,
                    spaceAfter=6
                )

                elements.append(Paragraph(para_text, para_style))
                elements.append(Spacer(1, 6))

        # Add page breaks if specified in template
        page_breaks = template_metadata.get("page_breaks", [])
        for page_break in page_breaks:
            elements.append(PageBreak())

        return elements

    def _is_markdown_table(self, text: str) -> bool:
        """Check if text is a markdown table"""
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        if len(lines) < 2:
            return False

        # Check for markdown table separator (| --- |)
        has_separator = any("|" in line and ("---" in line or "===" in line) for line in lines)
        has_pipes = all("|" in line for line in lines[:3])  # Check first few lines

        return has_separator and has_pipes

    def _parse_markdown_table(self, text: str) -> List[List[str]]:
        """Parse markdown table into list of rows"""
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        rows = []

        for line in lines:
            # Skip separator lines
            if "---" in line or "===" in line:
                continue

            # Split by pipe and clean up
            cells = [cell.strip() for cell in line.split("|")]
            # Remove empty cells at start/end
            cells = [c for c in cells if c]

            if cells:
                rows.append(cells)

        return rows if rows else None

    def _create_table_from_data(
        self,
        table_data: List[List[str]],
        template_metadata: Dict[str, Any]
    ) -> Table:
        """Create a ReportLab Table from parsed data"""
        if not table_data:
            return None

        # Get default font for table cells
        default_font = list(template_metadata.get("fonts", {}).keys())[0] if template_metadata.get("fonts") else "Helvetica"
        default_size = list(template_metadata.get("text_sizes", {}).keys())[0] if template_metadata.get("text_sizes") else 10

        # Convert data to Paragraph objects
        styles = getSampleStyleSheet()
        table_style = ParagraphStyle(
            name="Table",
            parent=styles["Normal"],
            fontName=self._registered_fonts.get(default_font, "Helvetica"),
            fontSize=default_size
        )

        formatted_data = []
        for row in table_data:
            formatted_row = []
            for cell in row:
                formatted_row.append(Paragraph(cell, table_style))
            formatted_data.append(formatted_row)

        # Create table
        table = Table(formatted_data)

        # Apply table style based on template
        table_style_config = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),  # Header row
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("FONTNAME", (0, 0), (-1, 0), self._registered_fonts.get(default_font, "Helvetica-Bold")),
            ("FONTSIZE", (0, 0), (-1, 0), default_size + 1),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
            ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
            ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ])

        table.setStyle(table_style_config)

        return table

    def generate_docx(
        self,
        content: str,
        template_metadata: Dict[str, Any],
        template_file_path: Optional[str] = None,
        output_name: str = "output.docx"
    ) -> str:
        """
        Generate a DOCX file from content using template metadata

        Args:
            content: The content to format
            template_metadata: Extracted template metadata
            template_file_path: Path to original template (for reference)
            output_name: Name for output DOCX

        Returns:
            Path to generated DOCX file
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX generation")

        doc = Document()

        # Apply section formatting
        sections = template_metadata.get("sections", [])
        if sections:
            section_info = sections[0]
            section = doc.sections[0]

            if section_info.get("page_width"):
                section.page_width = Pt(section_info["page_width"])
            if section_info.get("page_height"):
                section.page_height = Pt(section_info["page_height"])
            if section_info.get("left_margin"):
                section.left_margin = Pt(section_info["left_margin"])
            if section_info.get("right_margin"):
                section.right_margin = Pt(section_info["right_margin"])
            if section_info.get("top_margin"):
                section.top_margin = Pt(section_info["top_margin"])
            if section_info.get("bottom_margin"):
                section.bottom_margin = Pt(section_info["bottom_margin"])

        # Parse content
        paragraphs = content.split("\n\n")

        for para_text in paragraphs:
            if not para_text.strip():
                doc.add_paragraph()
                continue

            # Check if it's a header
            if para_text.strip().startswith("#"):
                level = len(para_text) - len(para_text.lstrip("#"))
                level = min(level, 6)

                text = para_text.lstrip("#").strip()
                doc.add_heading(text, level=level)

            # Check if it's a bullet list
            elif para_text.strip().startswith("-") or para_text.strip().startswith("*"):
                lines = para_text.split("\n")
                for line in lines:
                    if line.strip().startswith("-") or line.strip().startswith("*"):
                        text = line.lstrip("-*").strip()
                        doc.add_paragraph(text, style="List Bullet")

            # Regular paragraph
            else:
                para = doc.add_paragraph(para_text)

                # Apply formatting from template if available
                fonts = template_metadata.get("fonts", {})
                if fonts:
                    default_font = list(fonts.keys())[0]
                    for run in para.runs:
                        run.font.name = default_font.split(",")[0] if "," in default_font else default_font

        # Save document
        output_path = os.path.join(self.temp_dir, output_name)
        doc.save(output_path)

        return output_path
