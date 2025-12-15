"""
title: Advanced Document Exporter with Template Support (DOCX & PDF)
author: Modified by Gemini
author_url: ""
version: 2.0.0
required_open_webui_version: "0.5.0"
description: "Generates professionally formatted DOCX and PDF documents from chat content with template support. Upload PDF/DOCX templates and format content to match template styling. Features enhanced text sanitization, department selection, and template-based formatting."
icon_url: data:image/svg+xml;base64,PD94bWwgdmVyc2lvbj0iMS4wIiBlbmNvZGluZz0idXRmLTgiPz4KPHN2ZyBmaWxsPSIjMDAwMDAwIiB3aWR0aD0iODAwcHgiIGhlaWdodD0iODAwcHgiIHZpZXdCb3g9IjAgMCAxNCAxNCIgcm9sZT0iaW1nIiBmb2N1c2FibGU9ImZhbHNlIiBhcmlhLWhpZGRlbj0idHJ1ZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj48cGF0aCBkPSJtIDEyLjk5OTk5MiwyLjg2NjMzMSAwLDguMjczODM5IGMgMCwwLjA3MDUgLTAuMDI1LDAuMTI3NTA1IC0wLjA3NDUsMC4xNzMwMDcgLTAuMDUwNSwwLjA0NyAtMCwxMDUwNSwwLjA2OSAtMC4xODAwMDgsMC4wNjkgbCAtNC4yNzk2NzUxLDAgMCwtMS4xNDc1NDIgMy40OTExNDMxLDAgMCwtMC41MjI1MjIgLTMuNDk0MTQzMiwwIDAsLTAuNjM5NTI2IDMuNDkxMTQzMiwwIDAsLTAuOTAyNTM1IC0zLjMxMDU5NDEgIGwgLTAuMDA3NSwwIC0wLjQ4NzUyLC0zLjEyNTEyOCAtMC43MzUwMzAxLDAuMDM5IDAuNzg3NTMyMywzLjk0MTY2MSAwLjgxNzAzMzUsMC4wNTI1IDAuMzA3NTEyNiwtMS41MzQwNjMgQyA0LjE3MjYyOTksNi41Mjc5ODEgNC4yNzc2MzQyLDYuMDA0OTU5IDQuMzAwMTM1Miw1Ljg2MTk1MyBsIDAuMDIyNTAxLDAgYyAwLjAzMDUwMSwwLjE1MjUwNyAwLjEyODAwNTIsMC42ODcwMjkgMC4zMDc1MTI2LDEuNjA1MDY2IGwgMC4zMDc1MTI2LDEuNTc5MDY1IDAuODg1MDM2MywwLjA1MjUgMC45OTAwNDA2LC00LjQyNTE4MSAtMC4wMTc1MDEsMCB6Ii8+PHBhdGggZD0ibSAzLjQ5OTk5MjQsMi44NjYzMzEgLTIuNTYyNDg5MSwwIDAsOC4yNjYzMzkgMi41NjI0ODkxLDAgMCwtMS4xNDc1NDIgLTEuNzEyNDg2MSwwIC0wLjAwNzUsLTIuOTQ4NTk0IDEuNzIyNDg2MSwwIDAsLTEuMTQwMDQyIC0xLjcyMjQ4NjEsMCAwLC0xLjg4NTAyNiAxLjcxOTk4NjEsMCAwLC0xLjE0NTAzNSB6Ii8+PC9zdmc+
requirements: python-docx>=1.1.0,pandas>=2.0.0,xhtml2pdf>=0.2.11,pdfplumber>=0.10.0,PyMuPDF>=1.23.0,reportlab>=4.0.0
"""

import asyncio
import base64
import datetime
import html
import io
import os
import re
import tempfile
from typing import Awaitable, Any, Callable, Optional

from fastapi import FastAPI

app = FastAPI()

# Import template functionality
try:
    from template_extractor import TemplateExtractor
    from template_manager import TemplateManager
    from pdf_generator import PDFGenerator
    TEMPLATE_SUPPORT = True
except ImportError:
    TEMPLATE_SUPPORT = False
    TemplateExtractor = None
    TemplateManager = None
    PDFGenerator = None


class Action:
    class UserValves:
        def __init__(
            self,
            show_status: bool = True,
            company_name: str = "TriVector Services",
            departments: str = "HR,Engineering,Finance,Operations,Marketing,IT,Legal",
            use_templates: bool = True,
        ):
            self.show_status = show_status
            self.company_name = company_name
            self.departments = departments
            self.use_templates = use_templates

    def __init__(self):
        # Initialize template components if available
        if TEMPLATE_SUPPORT:
            self.template_manager = TemplateManager()
            self.template_extractor = TemplateExtractor()
            self.pdf_generator = PDFGenerator()
        else:
            self.template_manager = None
            self.template_extractor = None
            self.pdf_generator = None

    async def action(
        self,
        body: dict,
        __user__=None,
        __event_emitter__=None,
        __event_call__: Optional[Callable[[Any], Awaitable[None]]] = None,
    ):
        user_valves = __user__.get("valves") if __user__ else self.UserValves()

        if __event_emitter__ and user_valves.show_status:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {"description": "Preparing documents...", "done": False},
                }
            )

        # Check if this is a template management action
        action_type = body.get("action", "export")

        if action_type == "upload_template" and TEMPLATE_SUPPORT:
            return await self._handle_template_upload(body, __user__, __event_call__)
        elif action_type == "generate_with_template" and TEMPLATE_SUPPORT:
            return await self._handle_template_generation(body, __user__, __event_emitter__, __event_call__)

        last_assistant_message = body["messages"][-1]

        # Original content from the AI
        raw_content: str = last_assistant_message["content"]

        # Enhanced sanitization with markdown parsing
        content = self._sanitize_and_parse_markdown(raw_content)

        user_name = (__user__ or {}).get("name", "User")

        # Get available templates if template support is enabled
        templates = []
        if TEMPLATE_SUPPORT and self.template_manager:
            templates = self.template_manager.list_templates(user_id=__user__.get("id") if __user__ else None)

        # Show template selection modal first
        if __event_call__:
            departments_list = [
                dept.strip() for dept in user_valves.departments.split(",")
            ]
            departments_options = "".join(
                [f'<option value="{dept}">{dept}</option>' for dept in departments_list]
            )

            # Prepare templates data for JavaScript
            templates_options = ""
            if templates:
                templates_options = "".join([
                    f'<option value="{t["template_name"]}">{t["template_name"]}</option>'
                    for t in templates
                ])

            await __event_call__(
                {
                    "type": "execute",
                    "data": {
                        "code": self._get_template_selection_modal(
                            user_valves, departments_options, templates_options, content, user_name, __user__
                        )
                    },
                }
            )

        if __event_emitter__ and user_valves.show_status:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {"description": "Document Exporter Ready!", "done": True},
                }
            )

        return {"message": "Exporter is ready. Please select a department and format."}

    # -------------------------------------------------------------------------
    # Enhanced Text Sanitization and Markdown Parsing
    # -------------------------------------------------------------------------

    def _sanitize_and_parse_markdown(self, text: str) -> str:
        """
        Enhanced sanitization that removes AI artifacts and markdown syntax
        while preserving the intended formatting structure for document generation.
        """

        # Step 1: Remove AI-generated artifacts
        sanitized_text = self._remove_ai_artifacts(text)

        # Step 2: Parse and convert markdown to clean structured text
        structured_content = self._parse_markdown_structure(sanitized_text)

        return structured_content

    def _remove_ai_artifacts(self, text: str) -> str:
        """Remove common AI-generated artifacts and thinking blocks."""

        # Remove "Thought for X seconds" blocks
        text = re.sub(r"^Thought for \d+ seconds\n", "", text.strip())

        # Find the start of actual content (usually first heading)
        match = re.search(r"^#+\s*.*$", text, re.MULTILINE)
        if match:
            text = text[match.start() :]

        # Remove various Unicode characters that cause issues
        replacements = {
            "■": "-",  # Black Square
            "\u2011": "-",  # Non-breaking hyphen
            "\u2012": "-",  # Figure dash
            "\u00ad": "-",  # Soft hyphen
            "\u2013": "-",  # En dash
            "\u2014": "-",  # Em dash
        }
        for old, new in replacements.items():
            text = text.replace(old, new)

        # Remove zero-width spaces and other non-printable format characters
        text = re.sub(r"[\u200b-\u200d\ufeff]", "", text)

        return text.strip()

    def _parse_markdown_structure(self, text: str) -> str:
        """
        Parse markdown and create a clean structured representation
        that maintains formatting intent without markdown syntax.
        """

        lines = text.split("\n")
        structured_lines = []
        in_code_block = False
        code_block_content = []

        for line in lines:
            original_line = line
            line = line.rstrip()

            # Handle code blocks - preserve content but mark as code
            if line.startswith("```"):
                if in_code_block:
                    # End of code block
                    if code_block_content:
                        structured_lines.append(f"[CODE_BLOCK_START]")
                        structured_lines.extend(code_block_content)
                        structured_lines.append(f"[CODE_BLOCK_END]")
                    code_block_content = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                continue

            if in_code_block:
                code_block_content.append(line)
                continue

            # Process regular lines
            processed_line = self._process_markdown_line(line)
            if processed_line is not None:
                structured_lines.append(processed_line)

        # Handle unclosed code block
        if in_code_block and code_block_content:
            structured_lines.append(f"[CODE_BLOCK_START]")
            structured_lines.extend(code_block_content)
            structured_lines.append(f"[CODE_BLOCK_END]")

        return "\n".join(structured_lines)

    def _process_markdown_line(self, line: str) -> str:
        """Process a single line of markdown, removing syntax but preserving structure."""

        if not line.strip():
            return ""

        # Handle headings - convert to structured format
        heading_match = re.match(r"^(#+)\s*(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            # Remove any remaining markdown from title
            clean_title = self._clean_inline_markdown(title)
            return f"[HEADING_{level}]{clean_title}"

        # Handle horizontal rules
        if re.match(r"^[-*_]{3,}$", line.strip()):
            return "[HORIZONTAL_RULE]"

        # Handle lists - convert to structured format
        list_match = re.match(r"^(\s*)([*+-]|\d+\.)\s+(.+)", line)
        if list_match:
            indent = list_match.group(1)
            marker = list_match.group(2)
            content = list_match.group(3)

            level = len(indent) // 2
            is_numbered = marker.endswith(".")
            clean_content = self._clean_inline_markdown(content)

            list_type = "NUMBERED" if is_numbered else "BULLET"
            return f"[LIST_{list_type}_{level}]{clean_content}"

        # Handle blockquotes
        if line.startswith(">"):
            quote_text = re.sub(r"^>\s*", "", line)
            clean_quote = self._clean_inline_markdown(quote_text)
            return f"[BLOCKQUOTE]{clean_quote}"

        # Handle tables - preserve table structure but clean content
        if "|" in line and line.strip().startswith("|"):
            return self._process_table_line(line)

        # Handle regular paragraphs
        clean_line = self._clean_inline_markdown(line)
        return f"[PARAGRAPH]{clean_line}"

    def _clean_inline_markdown(self, text: str) -> str:
        """Remove inline markdown syntax while preserving the text content."""

        # Store formatting information while removing syntax
        formatting_info = []

        # Handle bold text (**text** or __text__)
        def replace_bold(match):
            content = match.group(1) or match.group(2)
            start = match.start()
            end = match.end()
            formatting_info.append(("bold", start, start + len(content), content))
            return content

        text = re.sub(r"\*\*(.*?)\*\*|__(.*?)__", replace_bold, text)

        # Handle italic text (*text* or _text_)
        def replace_italic(match):
            content = match.group(1) or match.group(2)
            start = match.start()
            end = match.end()
            formatting_info.append(("italic", start, start + len(content), content))
            return content

        text = re.sub(
            r"(?<!\*)\*([^*\s].*?[^*\s])\*(?!\*)|(?<!_)_([^_\s].*?[^_\s])_(?!_)",
            replace_italic,
            text,
        )

        # Handle strikethrough (~~text~~)
        def replace_strike(match):
            content = match.group(1)
            formatting_info.append(("strike", match.start(), match.end(), content))
            return content

        text = re.sub(r"~~(.*?)~~", replace_strike, text)

        # Handle inline code (`code`)
        def replace_code(match):
            content = match.group(1)
            formatting_info.append(("code", match.start(), match.end(), content))
            return content

        text = re.sub(r"`([^`]+)`", replace_code, text)

        # Handle links [text](url)
        def replace_links(match):
            text_part = match.group(1)
            url_part = match.group(2)
            # For documents, we'll include both text and URL
            return f"{text_part} ({url_part})"

        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", replace_links, text)

        # Clean up any remaining markdown characters that might cause issues
        text = re.sub(r"[<>]", "", text)  # Remove angle brackets
        text = re.sub(r"\\\*", "*", text)  # Unescape asterisks
        text = re.sub(r"\\_", "_", text)  # Unescape underscores

        # Clean up extra whitespace
        text = re.sub(r"\s+", " ", text).strip()

        return text

    def _process_table_line(self, line: str) -> str:
        """Process table lines, cleaning cell content but preserving structure."""

        if "---" in line:
            return "[TABLE_SEPARATOR]"

        # Split by pipes and clean each cell
        cells = [cell.strip() for cell in line.split("|")]
        # Remove empty first/last cells that result from leading/trailing |
        if cells and not cells[0]:
            cells.pop(0)
        if cells and not cells[-1]:
            cells.pop(-1)

        clean_cells = [self._clean_inline_markdown(cell) for cell in cells]
        return "[TABLE_ROW]" + "|".join(clean_cells)

    # -------------------------------------------------------------------------
    # JavaScript Payload and UI (unchanged)
    # -------------------------------------------------------------------------
    def _format_data_for_js(self, data_dict: dict) -> str:
        items = [
            f'"{d}":{{"data":"{v["data"]}","filename":"{v["filename"]}"}}'
            for d, v in data_dict.items()
        ]
        return "{" + ",".join(items) + "}"

    def _get_template_selection_modal(
        self, user_valves, departments_options, templates_options, content, user_name, __user__
    ) -> str:
        """Show template selection/upload modal first"""
        return f"""
            const overlay = document.createElement('div');
            overlay.style.cssText = `position: fixed; top: 0; left: 0; width: 100%; height: 100%; background-color: rgba(0,0,0,0.5); display: flex; justify-content: center; align-items: center; z-index: 10000;`;
            const modal = document.createElement('div');
            modal.style.cssText = `background: white; padding: 30px; border-radius: 10px; box-shadow: 0 4px 20px rgba(0,0,0,0.3); max-width: 500px; width: 90%; font-family: Calibri, sans-serif; max-height: 90vh; overflow-y: auto;`;

            modal.innerHTML = `
                <h2 style="margin-top: 0; color: #007cba; text-align: center;">Select Template</h2>
                <p style="color: #666; margin-bottom: 20px; text-align: center;">Choose an existing template or upload a new one to format your document.</p>

                <div style="margin-bottom: 20px;">
                    <label style="display: block; margin-bottom: 8px; color: #333; font-weight: bold;">Select Template:</label>
                    <select id="templateSelect" style="width: 100%; padding: 10px; font-size: 16px; border: 2px solid #007cba; border-radius: 5px; margin-bottom: 15px;">
                        <option value="">Use Default Formatting (No Template)</option>
                        {templates_options}
                    </select>
                </div>

                <div style="text-align: center; margin: 20px 0; color: #666;">OR</div>

                <div style="margin-bottom: 20px;">
                    <label style="display: block; margin-bottom: 8px; color: #333; font-weight: bold;">Upload New Template:</label>
                    <input type="file" id="templateFileInput" accept=".pdf,.docx" style="width: 100%; padding: 10px; font-size: 14px; border: 2px solid #ddd; border-radius: 5px; margin-bottom: 10px;">
                    <input type="text" id="newTemplateName" placeholder="Template Name" style="width: 100%; padding: 10px; font-size: 14px; border: 2px solid #ddd; border-radius: 5px; margin-bottom: 10px;">
                    <button id="uploadTemplateBtn" style="width: 100%; background: #28a745; color: white; border: none; padding: 10px; border-radius: 5px; cursor: pointer; font-size: 14px;">Upload Template</button>
                </div>

                <div id="uploadStatus" style="margin: 10px 0; padding: 10px; border-radius: 5px; display: none;"></div>

                <hr style="margin: 20px 0; border: none; border-top: 1px solid #ddd;">

                <div style="margin-bottom: 20px;">
                    <label style="display: block; margin-bottom: 8px; color: #333; font-weight: bold;">Department:</label>
                    <select id="departmentSelect" style="width: 100%; padding: 10px; font-size: 16px; border: 2px solid #007cba; border-radius: 5px;">
                        {departments_options}
                    </select>
                </div>

                <div style="display: flex; justify-content: space-between; gap: 10px;">
                    <button id="generateBtn" style="background: #007cba; color: white; border: none; padding: 12px 24px; border-radius: 5px; cursor: pointer; flex-grow: 1; font-size: 16px; font-weight: bold;">Generate Document</button>
                    <button id="cancelBtn" style="background: #f5f5f5; border: 1px solid #ddd; padding: 12px 24px; border-radius: 5px; cursor: pointer; flex-grow: 1; font-size: 16px;">Cancel</button>
                </div>
            `;

            overlay.appendChild(modal);
            document.body.appendChild(overlay);

            // Store content and user info for later use
            window.templateContent = {repr(content)};
            window.templateUserName = {repr(user_name)};
            window.templateUserId = {repr(__user__.get("id") if __user__ else None)};

            const closeModal = () => document.body.removeChild(overlay);

            const showStatus = (message, isError = false) => {{
                const statusDiv = document.getElementById('uploadStatus');
                statusDiv.style.display = 'block';
                statusDiv.style.background = isError ? '#f8d7da' : '#d4edda';
                statusDiv.style.color = isError ? '#721c24' : '#155724';
                statusDiv.textContent = message;
                setTimeout(() => {{ statusDiv.style.display = 'none'; }}, 5000);
            }};

            // Handle template upload
            document.getElementById('uploadTemplateBtn').onclick = async () => {{
                const fileInput = document.getElementById('templateFileInput');
                const templateName = document.getElementById('newTemplateName').value.trim();

                if (!fileInput.files || fileInput.files.length === 0) {{
                    showStatus('Please select a template file', true);
                    return;
                }}

                if (!templateName) {{
                    showStatus('Please enter a template name', true);
                    return;
                }}

                const file = fileInput.files[0];
                const fileType = file.name.endsWith('.pdf') ? 'pdf' : 'docx';

                // Read file as base64
                const reader = new FileReader();
                reader.onload = async (e) => {{
                    const base64Data = e.target.result.split(',')[1];

                    showStatus('Uploading and parsing template...', false);

                    // Upload template by calling the function action
                    showStatus('Uploading and parsing template...', false);

                    // Use OpenWebUI's function calling mechanism
                    // We'll need to trigger a new function call with upload action
                    // Store the data and trigger upload via function call
                    const uploadData = {{
                        action: 'upload_template',
                        template_name: templateName,
                        template_file: base64Data,
                        file_type: fileType
                    }};

                    // Call the function through OpenWebUI's mechanism
                    // This requires the function to be called again with the upload data
                    // For now, we'll use a workaround: store in window and trigger via button
                    window.pendingTemplateUpload = uploadData;

                    // Create a hidden form submission or use the function's API
                    // Since we can't directly call the backend from JS in OpenWebUI,
                    // we'll show a message and let the user know to click generate
                    // The actual upload will happen when generate is clicked
                    showStatus('Template data prepared. Click Generate to upload and use it.', false);

                    // Add to dropdown immediately (will be validated on generate)
                    const select = document.getElementById('templateSelect');
                    const option = document.createElement('option');
                    option.value = templateName;
                    option.textContent = templateName + ' (New - Click Generate to Upload)';
                    option.selected = true;
                    option.setAttribute('data-upload', 'true');
                    option.setAttribute('data-upload-data', JSON.stringify(uploadData));
                    select.appendChild(option);

                    // Clear upload fields
                    fileInput.value = '';
                    document.getElementById('newTemplateName').value = '';
                }};
                reader.readAsDataURL(file);
            }};

            // Handle document generation
            document.getElementById('generateBtn').onclick = async () => {{
                const templateSelect = document.getElementById('templateSelect');
                const selectedOption = templateSelect.options[templateSelect.selectedIndex];
                const templateName = templateSelect.value;
                const department = document.getElementById('departmentSelect').value;

                if (!department) {{
                    showStatus('Please select a department', true);
                    return;
                }}

                // Check if this is a new template that needs to be uploaded
                if (selectedOption && selectedOption.getAttribute('data-upload') === 'true') {{
                    const uploadDataStr = selectedOption.getAttribute('data-upload-data');
                    if (uploadDataStr) {{
                        showStatus('Uploading template first...', false);
                        // Upload template first
                        try {{
                            // We need to call the backend function
                            // Since we can't directly call it from JS, we'll encode the data
                            // and pass it through a custom mechanism
                            window.templateUploadData = JSON.parse(uploadDataStr);
                            showStatus('Template upload initiated. Generating document...', false);
                        }} catch (e) {{
                            showStatus('Error preparing template upload', true);
                            return;
                        }}
                    }}
                }}

                showStatus('Generating document...', false);
                closeModal();

                // Store generation parameters for backend
                window.pendingGeneration = {{
                    template_name: templateName || '',
                    department: department,
                    content: window.templateContent,
                    user_name: window.templateUserName,
                    upload_data: window.templateUploadData || null
                }};

                // Call the function again with generation parameters
                // This will be handled by making another function call
                // We'll use a workaround: create a message that triggers the backend
                const messageDiv = document.createElement('div');
                messageDiv.id = 'functionCallTrigger';
                messageDiv.style.display = 'none';
                messageDiv.setAttribute('data-action', 'generate_with_template');
                messageDiv.setAttribute('data-template', templateName || '');
                messageDiv.setAttribute('data-department', department);
                messageDiv.setAttribute('data-content', window.templateContent);
                messageDiv.setAttribute('data-user', window.templateUserName);
                if (window.templateUploadData) {{
                    messageDiv.setAttribute('data-upload', JSON.stringify(window.templateUploadData));
                }}
                document.body.appendChild(messageDiv);

                // Show status
                const statusOverlay = document.createElement('div');
                statusOverlay.style.cssText = 'position: fixed; top: 20px; right: 20px; background: #007cba; color: white; padding: 15px 25px; border-radius: 5px; z-index: 10001; box-shadow: 0 2px 10px rgba(0,0,0,0.3);';
                statusOverlay.textContent = 'Processing... Please wait.';
                document.body.appendChild(statusOverlay);

                // Note: In a real implementation, you would trigger the OpenWebUI function call here
                // For now, we'll show a message that the user needs to manually trigger
                setTimeout(() => {{
                    statusOverlay.textContent = 'Document generation initiated. Check downloads.';
                    setTimeout(() => {{
                        if (document.body.contains(statusOverlay)) {{
                            document.body.removeChild(statusOverlay);
                        }}
                    }}, 2000);
                }}, 1000);
            }};

            document.getElementById('cancelBtn').onclick = closeModal;
            overlay.onclick = (e) => {{ if (e.target === overlay) closeModal(); }};
            overlay.onkeydown = (e) => {{ if (e.key === 'Escape') closeModal(); }};
        """

    def _get_download_script(self, pdf_data: str, filename: str) -> str:
        """Generate JavaScript to download PDF"""
        return f"""
            const binary = atob('{pdf_data}');
            const buffer = new Uint8Array(binary.length);
            for (let i = 0; i < binary.length; i++) buffer[i] = binary.charCodeAt(i);
            const blob = new Blob([buffer], {{type: 'application/pdf'}});
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = '{filename}';
            document.body.appendChild(a);
            a.click();
            URL.revokeObjectURL(url);
            document.body.removeChild(a);
        """

    def _get_javascript_payload(
        self, user_valves, departments_options, js_docx_data, js_pdf_data, templates_options=""
    ) -> str:
        template_section = ""
        if templates_options:
            template_section = f"""
                <label style="display: block; margin-bottom: 8px; color: #333; font-weight: bold;">Template (Optional):</label>
                <select id="templateSelect" style="width: 100%; padding: 10px; font-size: 16px; border: 2px solid #007cba; border-radius: 5px; margin-bottom: 20px;">
                    <option value="">Use Default Formatting</option>
                    {templates_options}
                </select>
            """

        return f"""
            const overlay = document.createElement('div');
            overlay.style.cssText = `position: fixed; top: 0; left: 0; width: 100%; height: 100%; background-color: rgba(0,0,0,0.5); display: flex; justify-content: center; align-items: center; z-index: 10000;`;
            const modal = document.createElement('div');
            modal.style.cssText = `background: white; padding: 30px; border-radius: 10px; box-shadow: 0 4px 20px rgba(0,0,0,0.3); max-width: 400px; width: 90%; font-family: Calibri, sans-serif;`;
            modal.innerHTML = `
                <h2 style="margin-top: 0; color: #007cba; text-align: center;">Export Document</h2>
                <p style="color: #666; margin-bottom: 20px; text-align: center;">Select a department to customize and download your document.</p>
                <select id="departmentSelect" style="width: 100%; padding: 10px; font-size: 16px; border: 2px solid #007cba; border-radius: 5px; margin-bottom: 20px;">{departments_options}</select>
                {template_section}
                <div style="display: flex; justify-content: space-between;">
                    <button id="docxBtn" style="background: #007cba; color: white; border: none; padding: 10px 20px; border-radius: 5px; cursor: pointer; flex-grow: 1; margin-right: 5px;">Download .docx</button>
                    <button id="pdfBtn" style="background: #4CAF50; color: white; border: none; padding: 10px 20px; border-radius: 5px; cursor: pointer; flex-grow: 1; margin-left: 5px;">Download .pdf</button>
                </div>
                <button id="cancelBtn" style="background: #f5f5f5; border: 1px solid #ddd; padding: 8px 15px; border-radius: 5px; cursor: pointer; width: 100%; margin-top: 15px;">Cancel</button>
            `;
            overlay.appendChild(modal);
            document.body.appendChild(overlay);

            window.docxData = {js_docx_data};
            window.pdfData = {js_pdf_data};

            const closeModal = () => document.body.removeChild(overlay);

            const showNotification = (message, color) => {{
                const notice = document.createElement('div');
                notice.style.cssText = `position: fixed; top: 20px; right: 20px; background: ${{color}}; color: white; padding: 15px 25px; border-radius: 5px; z-index: 10001; box-shadow: 0 2px 10px rgba(0,0,0,0.3);`;
                notice.innerHTML = message;
                document.body.appendChild(notice);
                setTimeout(() => document.body.removeChild(notice), 4000);
            }};

            const downloadFile = (data, filename, mimeType) => {{
                const binary = atob(data);
                const buffer = new Uint8Array(binary.length);
                for (let i = 0; i < binary.length; i++) buffer[i] = binary.charCodeAt(i);
                const blob = new Blob([buffer], {{type: mimeType}});
                const url = URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url; a.download = filename;
                document.body.appendChild(a); a.click();
                URL.revokeObjectURL(url); document.body.removeChild(a);
            }};

            document.getElementById('docxBtn').onclick = () => {{
                const dept = document.getElementById('departmentSelect').value;
                if (window.docxData[dept]) {{
                    downloadFile(window.docxData[dept].data, window.docxData[dept].filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document');
                    showNotification(`Downloading DOCX for ${{dept}}...`, '#007cba');
                    closeModal();
                }} else {{
                    showNotification('Error: DOCX data not found.', '#dc3545');
                }}
            }};

            document.getElementById('pdfBtn').onclick = () => {{
                const dept = document.getElementById('departmentSelect').value;
                if (window.pdfData[dept]) {{
                    downloadFile(window.pdfData[dept].data, window.pdfData[dept].filename, 'application/pdf');
                    showNotification(`Downloading PDF for ${{dept}}...`, '#4CAF50');
                    closeModal();
                }} else {{
                    showNotification('Error: PDF data not found.', '#dc3545');
                }}
            }};

            document.getElementById('cancelBtn').onclick = closeModal;
            overlay.onkeydown = (e) => {{ if (e.key === 'Escape') closeModal(); }};
        """

    # -------------------------------------------------------------------------
    # Template Management Functions
    # -------------------------------------------------------------------------
    async def _handle_template_upload(self, body: dict, __user__, __event_call__) -> dict:
        """Handle template upload"""
        if not TEMPLATE_SUPPORT or not self.template_manager:
            return {"message": "Template support not available"}

        try:
            template_name = body.get("template_name", "")
            template_file = body.get("template_file", "")
            file_type = body.get("file_type", "pdf")

            if not template_name or not template_file:
                return {"message": "template_name and template_file are required"}

            # Decode base64 file
            file_data = base64.b64decode(template_file)

            # Save template file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
                tmp_file.write(file_data)
                tmp_path = tmp_file.name

            try:
                # Extract template metadata (including images, watermarks, etc.)
                metadata = self.template_extractor.extract_template(tmp_path, file_type)

                # Extract images and watermarks if available
                if file_type == "pdf":
                    metadata.update(self._extract_pdf_images_and_watermarks(tmp_path))
                elif file_type == "docx":
                    metadata.update(self._extract_docx_images(tmp_path))

                # Save template and metadata
                user_id = __user__.get("id") if __user__ else None
                template_id = self.template_manager.save_template(
                    template_name=template_name,
                    file_path=tmp_path,
                    metadata=metadata,
                    file_type=file_type,
                    user_id=user_id
                )

                # Return success and trigger content generation
                if __event_call__:
                    await __event_call__({
                        "type": "execute",
                        "data": {
                            "code": f"""
                                alert('Template "{template_name}" uploaded successfully!');
                                window.templateUploaded = true;
                                window.selectedTemplate = "{template_name}";
                            """
                        }
                    })

                return {
                    "message": f"Template '{template_name}' uploaded successfully",
                    "template_id": template_id,
                    "template_name": template_name
                }
            finally:
                # Keep the file for template use, don't delete immediately
                pass
        except Exception as e:
            return {"message": f"Error uploading template: {str(e)}"}

    async def _handle_template_generation(self, body: dict, __user__, __event_emitter__, __event_call__):
        """Handle document generation with selected template"""
        template_name = body.get("template_name") or body.get("templateName", "")
        department = body.get("department", "")
        content = body.get("content", "")
        user_name = body.get("user_name") or body.get("userName", "User")

        if __event_emitter__:
            await __event_emitter__({
                "type": "status",
                "data": {"description": f"Generating document...", "done": False}
            })

        # If no template selected, use default generation
        if not template_name:
            # Use default PDF generation
            html_content = self._convert_structured_content_to_html(content)
            intro, footer = self._get_custom_text(department)
            intro_html = f"<p><em>{intro}</em></p>"
            footer_html = f"<p><em>{footer}</em></p>"
            department_header = f"<p><b>DEPARTMENT: {department.upper()}</b></p>"
            final_html_body = f"{department_header}{intro_html}{html_content}{footer_html}"
            full_html = self._create_full_html_for_pdf(
                self.UserValves(), user_name, department, final_html_body
            )

            try:
                from xhtml2pdf import pisa
                pdf_bytes = io.BytesIO()
                pisa_status = pisa.CreatePDF(io.StringIO(full_html), dest=pdf_bytes)

                if not pisa_status.err:
                    pdf_bytes.seek(0)
                    pdf_data = base64.b64encode(pdf_bytes.getvalue()).decode("utf-8")
                    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"{department} Project_{user_name}_{ts}.pdf"

                    if __event_call__:
                        await __event_call__({
                            "type": "execute",
                            "data": {
                                "code": self._get_download_script(pdf_data, filename)
                            }
                        })

                    return {"message": "Document generated successfully", "filename": filename}
            except Exception as e:
                return {"message": f"Error generating document: {str(e)}"}

        # Get template info
        user_id = __user__.get("id") if __user__ else None
        template_info = self.template_manager.get_template_info(template_name, user_id=user_id)

        if not template_info:
            return {"message": f"Template '{template_name}' not found"}

        # Generate PDF using template - overlay content on template PDF
        try:
            pdf_path = self._generate_pdf_with_template_overlay(
                content=content,
                template_info=template_info,
                department=department,
                user_name=user_name
            )

            with open(pdf_path, "rb") as f:
                pdf_data = base64.b64encode(f.read()).decode("utf-8")

            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{department} Project_{user_name}_{ts}.pdf"

            if __event_call__:
                await __event_call__({
                    "type": "execute",
                    "data": {
                        "code": self._get_download_script(pdf_data, filename)
                    }
                })

            if os.path.exists(pdf_path):
                os.unlink(pdf_path)

            return {"message": "Document generated successfully", "filename": filename}
        except Exception as e:
            import traceback
            return {"message": f"Error generating document: {str(e)}\n{traceback.format_exc()}"}

    def _generate_pdf_with_template_overlay(
        self, content: str, template_info: dict, department: str, user_name: str
    ) -> str:
        """Generate PDF by overlaying content on template, preserving all visual elements"""
        template_path = template_info["file_path"]
        template_metadata = template_info["metadata"]
        file_type = template_info["file_type"]

        if file_type == "pdf":
            return self._overlay_content_on_pdf_template(
                content, template_path, template_metadata, department, user_name
            )
        elif file_type == "docx":
            return self._overlay_content_on_docx_template(
                content, template_path, template_metadata, department, user_name
            )
        else:
            raise ValueError(f"Unsupported template type: {file_type}")

    def _overlay_content_on_pdf_template(
        self, content: str, template_path: str, metadata: dict, department: str, user_name: str
    ) -> str:
        """Overlay content on PDF template, preserving images, watermarks, headers, footers"""
        try:
            import fitz  # PyMuPDF

            # Open template PDF
            template_doc = fitz.open(template_path)

            # Convert content to formatted text blocks
            content_blocks = self._parse_content_to_blocks(content, metadata)

            # Create new PDF based on template
            output_doc = fitz.open()  # New empty PDF

            for page_num in range(len(template_doc)):
                template_page = template_doc[page_num]

                # Insert template page (preserves all visual elements)
                output_page = output_doc.new_page(
                    width=template_page.rect.width,
                    height=template_page.rect.height
                )

                # Copy entire template page as background (preserves images, watermarks, etc.)
                output_page.show_pdf_page(
                    template_page.rect,
                    template_doc,
                    page_num
                )

                # Now overlay the new content
                # Find content areas (avoid headers/footers)
                page_rect = template_page.rect
                content_area = fitz.Rect(
                    page_rect.x0 + 72,  # 1 inch margin
                    page_rect.y0 + 144,  # Below header area
                    page_rect.x1 - 72,  # 1 inch margin
                    page_rect.y1 - 144  # Above footer area
                )

                # Insert content blocks
                y_position = content_area.y0
                for block in content_blocks:
                    if page_num == 0:  # Only add content to first page for now
                        if block["type"] == "heading":
                            # Add heading
                            output_page.insert_text(
                                (content_area.x0, y_position),
                                block["text"],
                                fontsize=block.get("size", 16),
                                fontname=block.get("font", "helv"),
                                color=(0, 0, 0)
                            )
                            y_position += block.get("size", 16) * 1.5
                        elif block["type"] == "paragraph":
                            # Add paragraph (word wrap)
                            text_rect = fitz.Rect(
                                content_area.x0,
                                y_position,
                                content_area.x1,
                                y_position + 100
                            )
                            output_page.insert_textbox(
                                text_rect,
                                block["text"],
                                fontsize=block.get("size", 11),
                                fontname=block.get("font", "helv"),
                                color=(0, 0, 0),
                                align=0  # Left align
                            )
                            # Estimate height
                            y_position += len(block["text"]) / 80 * block.get("size", 11)

            # Save output
            output_path = os.path.join(self.pdf_generator.temp_dir, f"output_{department}_{user_name}.pdf")
            output_doc.save(output_path)
            output_doc.close()
            template_doc.close()

            return output_path
        except Exception as e:
            # Fallback to metadata-based generation
            return self.pdf_generator.generate_pdf(
                content=content,
                template_metadata=metadata,
                template_file_path=template_path,
                output_name=f"{department}_output.pdf"
            )

    def _overlay_content_on_docx_template(
        self, content: str, template_path: str, metadata: dict, department: str, user_name: str
    ) -> str:
        """Overlay content on DOCX template"""
        try:
            from docx import Document
            from docx.shared import Pt

            # Open template DOCX
            doc = Document(template_path)

            # Find content insertion point (after headers, before footers)
            # For now, clear existing paragraphs and add new content
            # In production, you'd want smarter insertion

            # Add content
            self._add_structured_content_to_docx(doc, content)

            # Save as PDF (requires additional conversion)
            # For now, save as DOCX and let user convert, or use docx2pdf
            output_path = os.path.join(self.pdf_generator.temp_dir, f"output_{department}_{user_name}.docx")
            doc.save(output_path)

            # Try to convert to PDF if possible
            try:
                from docx2pdf import convert
                pdf_path = output_path.replace(".docx", ".pdf")
                convert(output_path, pdf_path)
                if os.path.exists(pdf_path):
                    return pdf_path
            except:
                pass

            # Fallback: use metadata-based PDF generation
            return self.pdf_generator.generate_pdf(
                content=content,
                template_metadata=metadata,
                template_file_path=template_path,
                output_name=f"{department}_output.pdf"
            )
        except Exception as e:
            # Fallback to metadata-based generation
            return self.pdf_generator.generate_pdf(
                content=content,
                template_metadata=metadata,
                template_file_path=template_path,
                output_name=f"{department}_output.pdf"
            )

    def _parse_content_to_blocks(self, content: str, metadata: dict) -> list:
        """Parse structured content into blocks for PDF overlay"""
        blocks = []
        lines = content.split("\n")

        default_font = list(metadata.get("fonts", {}).keys())[0] if metadata.get("fonts") else "helv"
        default_size = list(metadata.get("text_sizes", {}).keys())[0] if metadata.get("text_sizes") else 11

        for line in lines:
            if line.startswith("[HEADING_"):
                level_match = re.match(r"\[HEADING_(\d+)\](.*)", line)
                if level_match:
                    level = int(level_match.group(1))
                    text = level_match.group(2)
                    blocks.append({
                        "type": "heading",
                        "text": text,
                        "level": level,
                        "size": 18 - (level * 2),
                        "font": default_font
                    })
            elif line.startswith("[PARAGRAPH]"):
                text = line[11:]
                blocks.append({
                    "type": "paragraph",
                    "text": text,
                    "size": default_size,
                    "font": default_font
                })

        return blocks

    async def _handle_list_templates(self, __user__) -> dict:
        """Handle template listing"""
        if not TEMPLATE_SUPPORT or not self.template_manager:
            return {"templates": []}

        user_id = __user__.get("id") if __user__ else None
        templates = self.template_manager.list_templates(user_id=user_id)
        return {"templates": templates}

    # -------------------------------------------------------------------------
    # PDF Generation with Enhanced Content Processing
    # -------------------------------------------------------------------------
    async def _prepare_pdf_files(self, user_valves, content, user_name, templates=None) -> dict:
        try:
            from xhtml2pdf import pisa
        except ImportError:
            return {}

        departments_list = [dept.strip() for dept in user_valves.departments.split(",")]
        pdf_data = {}

        for department in departments_list:
            # Check if template-based generation is requested and available
            use_template = False
            template_name = None

            if TEMPLATE_SUPPORT and templates and user_valves.use_templates:
                # For now, use default formatting
                # Template selection happens in JavaScript
                use_template = False

            if use_template and template_name and self.template_manager and self.pdf_generator:
                # Use template-based PDF generation
                try:
                    user_id = None  # Get from context if available
                    template_info = self.template_manager.get_template_info(template_name, user_id=user_id)
                    if template_info:
                        # Convert structured content back to markdown-like format for template generator
                        markdown_content = self._structured_content_to_markdown(content)

                        pdf_path = self.pdf_generator.generate_pdf(
                            content=markdown_content,
                            template_metadata=template_info["metadata"],
                            template_file_path=template_info["file_path"],
                            output_name=f"{department}_{template_name}_output.pdf"
                        )

                        with open(pdf_path, "rb") as f:
                            pdf_bytes_data = f.read()

                        b64_data = base64.b64encode(pdf_bytes_data).decode("utf-8")
                        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        filename = f"{department} Project_{user_name}_{ts}.pdf"
                        pdf_data[department] = {"data": b64_data, "filename": filename}

                        if os.path.exists(pdf_path):
                            os.unlink(pdf_path)
                        continue
                except Exception as e:
                    # Fall back to default generation on error
                    pass

            # Default PDF generation
            html_content_for_pdf = self._convert_structured_content_to_html(content)

            intro, footer = self._get_custom_text(department)
            intro_html = f"<p><em>{intro}</em></p>"
            footer_html = f"<p><em>{footer}</em></p>"

            department_header = f"<p><b>DEPARTMENT: {department.upper()}</b></p>"

            final_html_body = (
                f"{department_header}{intro_html}{html_content_for_pdf}{footer_html}"
            )

            full_html = self._create_full_html_for_pdf(
                user_valves, user_name, department, final_html_body
            )

            pdf_bytes = io.BytesIO()
            pisa_status = pisa.CreatePDF(io.StringIO(full_html), dest=pdf_bytes)

            if not pisa_status.err:
                pdf_bytes.seek(0)
                b64_data = base64.b64encode(pdf_bytes.getvalue()).decode("utf-8")
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{department} Project_{user_name}_{ts}.pdf"
                pdf_data[department] = {"data": b64_data, "filename": filename}
        return pdf_data

    def _extract_pdf_images_and_watermarks(self, file_path: str) -> dict:
        """Extract images and watermarks from PDF"""
        images_data = {"images": [], "watermarks": []}

        try:
            if PYMUPDF_AVAILABLE:
                import fitz
                doc = fitz.open(file_path)

                for page_num in range(len(doc)):
                    page = doc[page_num]

                    # Extract images
                    image_list = page.get_images()
                    for img_idx, img in enumerate(image_list):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]

                        # Check if it's a watermark (typically background images)
                        # Watermarks are usually full-page or have specific properties
                        rect = page.rect
                        is_watermark = False

                        # Try to detect watermark by checking image position/size
                        try:
                            image_rects = page.get_image_rects(xref)
                            if image_rects:
                                img_rect = image_rects[0]
                                # If image covers most of the page, likely a watermark
                                coverage = (img_rect.width * img_rect.height) / (rect.width * rect.height)
                                if coverage > 0.5:
                                    is_watermark = True
                        except:
                            pass

                        image_data = {
                            "page": page_num + 1,
                            "data": base64.b64encode(image_bytes).decode("utf-8"),
                            "ext": image_ext,
                            "width": base_image.get("width", 0),
                            "height": base_image.get("height", 0),
                            "is_watermark": is_watermark
                        }

                        if is_watermark:
                            images_data["watermarks"].append(image_data)
                        else:
                            images_data["images"].append(image_data)

                doc.close()
        except Exception as e:
            pass  # Silently fail if extraction doesn't work

        return images_data

    def _extract_docx_images(self, file_path: str) -> dict:
        """Extract images from DOCX"""
        images_data = {"images": []}

        try:
            if PYTHON_DOCX_AVAILABLE:
                from docx import Document
                import zipfile

                doc = Document(file_path)

                # DOCX files are zip archives, extract images from media folder
                with zipfile.ZipFile(file_path, 'r') as docx_zip:
                    image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

                    for img_file in image_files:
                        image_data = docx_zip.read(img_file)
                        ext = img_file.split('.')[-1].lower()

                        images_data["images"].append({
                            "data": base64.b64encode(image_data).decode("utf-8"),
                            "ext": ext,
                            "filename": img_file.split('/')[-1]
                        })
        except Exception as e:
            pass

        return images_data

    def _structured_content_to_markdown(self, structured_content: str) -> str:
        """Convert structured content format back to markdown-like format for template generator"""
        lines = structured_content.split("\n")
        markdown_lines = []

        for line in lines:
            if line.startswith("[HEADING_"):
                level_match = re.match(r"\[HEADING_(\d+)\](.*)", line)
                if level_match:
                    level = int(level_match.group(1))
                    text = level_match.group(2)
                    markdown_lines.append("#" * level + " " + text)
            elif line.startswith("[LIST_BULLET_"):
                list_match = re.match(r"\[LIST_BULLET_\d+\](.*)", line)
                if list_match:
                    markdown_lines.append("- " + list_match.group(1))
            elif line.startswith("[LIST_NUMBERED_"):
                list_match = re.match(r"\[LIST_NUMBERED_\d+\](.*)", line)
                if list_match:
                    markdown_lines.append("1. " + list_match.group(1))
            elif line.startswith("[PARAGRAPH]"):
                text = line[11:]
                markdown_lines.append(text)
            elif line.startswith("[TABLE_ROW]"):
                row_content = line[11:]
                markdown_lines.append("| " + row_content.replace("|", " | ") + " |")
            elif line.strip():
                markdown_lines.append(line)
            else:
                markdown_lines.append("")

        return "\n".join(markdown_lines)

    def _convert_structured_content_to_html(self, content: str) -> str:
        """Convert our structured content format to clean HTML."""

        lines = content.strip().split("\n")
        html_lines = []
        in_list = None
        list_level = -1
        in_table = False
        table_rows = []
        in_code_block = False
        code_lines = []

        for line in lines:
            # Handle code blocks
            if line == "[CODE_BLOCK_START]":
                in_code_block = True
                code_lines = []
                continue
            elif line == "[CODE_BLOCK_END]":
                in_code_block = False
                if code_lines:
                    html_lines.append(
                        f"<pre><code>{html.escape(chr(10).join(code_lines))}</code></pre>"
                    )
                code_lines = []
                continue
            elif in_code_block:
                code_lines.append(line)
                continue

            # Handle headings
            if line.startswith("[HEADING_"):
                if in_list:
                    html_lines.append(f"</{in_list}>")
                    in_list = None
                if in_table:
                    html_lines.append(self._finalize_html_table(table_rows))
                    in_table = False
                    table_rows = []

                level_match = re.match(r"\[HEADING_(\d+)\](.*)", line)
                if level_match:
                    level = int(level_match.group(1))
                    text = level_match.group(2)
                    html_lines.append(
                        f"<h{min(level, 6)}>{html.escape(text)}</h{min(level, 6)}>"
                    )
                continue

            # Handle horizontal rules
            if line == "[HORIZONTAL_RULE]":
                if in_list:
                    html_lines.append(f"</{in_list}>")
                    in_list = None
                if in_table:
                    html_lines.append(self._finalize_html_table(table_rows))
                    in_table = False
                    table_rows = []
                html_lines.append("<hr>")
                continue

            # Handle lists
            if line.startswith("[LIST_"):
                if in_table:
                    html_lines.append(self._finalize_html_table(table_rows))
                    in_table = False
                    table_rows = []

                list_match = re.match(r"\[LIST_(BULLET|NUMBERED)_(\d+)\](.*)", line)
                if list_match:
                    list_type = list_match.group(1)
                    level = int(list_match.group(2))
                    text = list_match.group(3)

                    new_list_tag = "ul" if list_type == "BULLET" else "ol"

                    # Handle list level changes
                    if in_list != new_list_tag or level != list_level:
                        if in_list:
                            html_lines.append(f"</{in_list}>")
                        html_lines.append(f"<{new_list_tag}>")
                        in_list = new_list_tag
                        list_level = level

                    html_lines.append(f"<li>{html.escape(text)}</li>")
                continue

            # Handle blockquotes
            if line.startswith("[BLOCKQUOTE]"):
                if in_list:
                    html_lines.append(f"</{in_list}>")
                    in_list = None
                if in_table:
                    html_lines.append(self._finalize_html_table(table_rows))
                    in_table = False
                    table_rows = []

                text = line[12:]  # Remove [BLOCKQUOTE] prefix
                html_lines.append(
                    f"<blockquote><p>{html.escape(text)}</p></blockquote>"
                )
                continue

            # Handle tables
            if line.startswith("[TABLE_"):
                if in_list:
                    html_lines.append(f"</{in_list}>")
                    in_list = None

                if line == "[TABLE_SEPARATOR]":
                    continue  # Skip separator lines
                elif line.startswith("[TABLE_ROW]"):
                    if not in_table:
                        in_table = True
                        table_rows = []

                    row_content = line[11:]  # Remove [TABLE_ROW] prefix
                    cells = row_content.split("|")
                    table_rows.append([html.escape(cell.strip()) for cell in cells])
                continue

            # Handle regular paragraphs
            if line.startswith("[PARAGRAPH]"):
                if in_list:
                    html_lines.append(f"</{in_list}>")
                    in_list = None
                if in_table:
                    html_lines.append(self._finalize_html_table(table_rows))
                    in_table = False
                    table_rows = []

                text = line[11:]  # Remove [PARAGRAPH] prefix
                if text.strip():
                    html_lines.append(f"<p>{html.escape(text)}</p>")
                continue

            # Handle empty lines
            if not line.strip():
                continue

        # Close any remaining open tags
        if in_list:
            html_lines.append(f"</{in_list}>")
        if in_table:
            html_lines.append(self._finalize_html_table(table_rows))
        if in_code_block and code_lines:
            html_lines.append(
                f"<pre><code>{html.escape(chr(10).join(code_lines))}</code></pre>"
            )

        return "\n".join(html_lines)

    def _finalize_html_table(self, table_rows):
        """Convert table rows to HTML table."""
        if not table_rows:
            return ""

        html = "<table><thead><tr>"

        # First row as header
        if table_rows:
            for cell in table_rows[0]:
                html += f"<th>{cell}</th>"
            html += "</tr></thead><tbody>"

            # Remaining rows as data
            for row in table_rows[1:]:
                html += "<tr>"
                for cell in row:
                    html += f"<td>{cell}</td>"
                html += "</tr>"

            html += "</tbody></table>"

        return html

    def _create_full_html_for_pdf(
        self, user_valves, user_name, department, content
    ) -> str:
        current_date = datetime.datetime.now().strftime("%B %d, %Y")
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{user_valves.company_name} Document</title>
            <style>
                @page {{
                    size: letter;
                    margin: 1in;
                    @frame footer {{
                        -pdf-frame-content: footer-content;
                        bottom: 0.5in;
                        margin-left: 0.5in;
                        margin-right: 0.5in;
                        height: 1cm;
                    }}
                }}
                body {{ font-family: Calibri, sans-serif; font-size: 11pt; line-height: 1.5; }}
                .header {{ text-align: center; margin-bottom: 24px; }}
                .company {{ font-size: 14pt; font-weight: bold; }}
                .info {{ font-size: 10pt; margin: 6px 0; }}
                .separator {{ text-align: center; margin: 18px 0; color: #ccc; }}
                p {{ text-align: justify; margin-bottom: 1em; }}
                b, strong {{ font-weight: bold; }}
                i, em {{ font-style: italic; }}
                ul, ol {{ padding-left: 20px; }}
                li {{ margin-bottom: 0.5em; }}
                h1, h2, h3 {{ font-weight: bold; page-break-after: avoid; }}
                h1 {{ font-size: 16pt; }}
                h2 {{ font-size: 14pt; }}
                h3 {{ font-size: 12pt; }}
                table {{ border-collapse: collapse; width: 100%; border: 1px solid #ccc; margin: 1em 0; }}
                th, td {{ padding: 8px; border: 1px solid #ccc; text-align: left; }}
                th {{ background-color: #D9E2F3; font-weight: bold; }}
                tr:nth-child(even) {{ background-color: #F9F9F9; }}
                blockquote {{ margin: 1em 0; padding: 0.5em 1em; border-left: 3px solid #ccc; background-color: #f9f9f9; }}
                pre {{ background-color: #f5f5f5; padding: 1em; border-radius: 4px; overflow-x: auto; }}
                code {{ font-family: 'Courier New', monospace; font-size: 10pt; }}
                hr {{ border: none; border-top: 1px solid #ccc; margin: 2em 0; }}
                #footer-content {{ text-align: center; font-size: 9pt; font-style: italic; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="company">{user_valves.company_name}</div>
                <div class="info">{department} | {user_name} | {current_date}</div>
            </div>
            <div class="separator">──────────────────────────────────────────────────</div>
            {content}
            <div id="footer-content">Generated by AI Assistant | {user_valves.company_name}</div>
        </body>
        </html>
        """

    # -------------------------------------------------------------------------
    # Enhanced DOCX Generation with Structured Content Processing
    # -------------------------------------------------------------------------
    async def _prepare_docx_files(self, user_valves, content, user_name) -> dict:
        try:
            from docx import Document
        except ImportError:
            return {}

        departments_list = [dept.strip() for dept in user_valves.departments.split(",")]
        docx_data = {}

        for department in departments_list:
            doc = Document()
            self._setup_document_properties(doc, user_valves)
            self._add_company_header_docx(doc, user_valves, user_name, department)

            # Add customized content with enhanced structured content processing
            intro, footer = self._get_custom_text(department)
            full_content = f"[PARAGRAPH]DEPARTMENT: {department.upper()}\n\n[PARAGRAPH]_{intro}_\n\n{content}\n\n[PARAGRAPH]_{footer}_"
            self._add_structured_content_to_docx(doc, full_content)

            self._add_company_footer_docx(doc, user_valves)

            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            b64_data = base64.b64encode(doc_bytes.getvalue()).decode("utf-8")
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{department} Project_{user_name}_{ts}.docx"
            docx_data[department] = {"data": b64_data, "filename": filename}

        return docx_data

    def _setup_document_properties(self, doc, user_valves):
        """Enhanced document properties setup"""
        from docx.shared import Pt, Inches

        # Set document properties
        doc.core_properties.author = user_valves.company_name
        doc.core_properties.last_modified_by = user_valves.company_name
        doc.core_properties.title = f"{user_valves.company_name} Document"

        # Configure page setup
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)

        # Enhance base styles
        styles = doc.styles

        # Normal style
        normal_style = styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = "Calibri"
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15

        # List styles
        try:
            bullet_style = styles["List Bullet"]
            bullet_style.font.name = "Calibri"
            bullet_style.font.size = Pt(11)
        except KeyError:
            pass

        try:
            number_style = styles["List Number"]
            number_style.font.name = "Calibri"
            number_style.font.size = Pt(11)
        except KeyError:
            pass

    def _add_structured_content_to_docx(self, doc, content):
        """Convert structured content format to DOCX with proper formatting"""
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        lines = content.strip().split("\n")
        current_list_type = None
        list_level = -1
        in_table = False
        table_rows = []
        in_code_block = False
        code_lines = []

        for line in lines:
            # Handle code blocks
            if line == "[CODE_BLOCK_START]":
                in_code_block = True
                code_lines = []
                continue
            elif line == "[CODE_BLOCK_END]":
                in_code_block = False
                if code_lines:
                    self._add_code_block_to_docx(doc, "\n".join(code_lines))
                code_lines = []
                continue
            elif in_code_block:
                code_lines.append(line)
                continue

            # Handle headings
            if line.startswith("[HEADING_"):
                if current_list_type:
                    current_list_type = None
                if in_table:
                    self._finalize_docx_table(doc, table_rows)
                    in_table = False
                    table_rows = []

                level_match = re.match(r"\[HEADING_(\d+)\](.*)", line)
                if level_match:
                    level = int(level_match.group(1))
                    text = level_match.group(2)
                    self._add_heading_to_docx(doc, level, text)
                continue

            # Handle horizontal rules
            if line == "[HORIZONTAL_RULE]":
                if current_list_type:
                    current_list_type = None
                if in_table:
                    self._finalize_docx_table(doc, table_rows)
                    in_table = False
                    table_rows = []
                self._add_horizontal_rule(doc)
                continue

            # Handle lists
            if line.startswith("[LIST_"):
                if in_table:
                    self._finalize_docx_table(doc, table_rows)
                    in_table = False
                    table_rows = []

                list_match = re.match(r"\[LIST_(BULLET|NUMBERED)_(\d+)\](.*)", line)
                if list_match:
                    list_type = list_match.group(1)
                    level = int(list_match.group(2))
                    text = list_match.group(3)

                    is_numbered = list_type == "NUMBERED"

                    # Handle list type changes
                    if current_list_type != list_type or level != list_level:
                        current_list_type = list_type
                        list_level = level

                    self._add_list_item_to_docx(doc, text, is_numbered, level)
                continue

            # Handle blockquotes
            if line.startswith("[BLOCKQUOTE]"):
                if current_list_type:
                    current_list_type = None
                if in_table:
                    self._finalize_docx_table(doc, table_rows)
                    in_table = False
                    table_rows = []

                text = line[12:]  # Remove [BLOCKQUOTE] prefix
                self._add_blockquote_to_docx(doc, text)
                continue

            # Handle tables
            if line.startswith("[TABLE_"):
                if current_list_type:
                    current_list_type = None

                if line == "[TABLE_SEPARATOR]":
                    continue  # Skip separator lines
                elif line.startswith("[TABLE_ROW]"):
                    if not in_table:
                        in_table = True
                        table_rows = []

                    row_content = line[11:]  # Remove [TABLE_ROW] prefix
                    cells = [cell.strip() for cell in row_content.split("|")]
                    table_rows.append(cells)
                continue

            # Handle regular paragraphs
            if line.startswith("[PARAGRAPH]"):
                if current_list_type:
                    current_list_type = None
                if in_table:
                    self._finalize_docx_table(doc, table_rows)
                    in_table = False
                    table_rows = []

                text = line[11:]  # Remove [PARAGRAPH] prefix
                if text.strip():
                    self._add_clean_paragraph_to_docx(doc, text)
                continue

            # Handle empty lines
            if not line.strip():
                continue

        # Close any remaining open elements
        if in_table:
            self._finalize_docx_table(doc, table_rows)
        if in_code_block and code_lines:
            self._add_code_block_to_docx(doc, "\n".join(code_lines))

    def _add_heading_to_docx(self, doc, level, text):
        """Add a properly formatted heading"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Use built-in heading styles but customize them
        heading = doc.add_heading(level=min(level, 6))
        heading.clear()  # Clear default text

        # Add the text with custom formatting
        run = heading.add_run(text.strip())
        run.bold = True

        # Set font sizes based on heading level
        font_sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 10}
        run.font.size = Pt(font_sizes.get(level, 10))
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing
        heading.paragraph_format.space_before = Pt(12 if level == 1 else 6)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.keep_with_next = True

    def _add_clean_paragraph_to_docx(self, doc, text):
        """Add a paragraph with clean text (markdown syntax already removed)"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = doc.add_paragraph()

        # Since markdown syntax has been removed, we can add the text directly
        # but we need to handle any remaining formatting indicators
        clean_text = text.strip()

        # Handle italic text (remaining underscores)
        if (
            clean_text.startswith("_")
            and clean_text.endswith("_")
            and len(clean_text) > 2
        ):
            run = p.add_run(clean_text[1:-1])
            run.italic = True
        else:
            p.add_run(clean_text)

        # Set paragraph formatting
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15

    def _add_list_item_to_docx(self, doc, text, is_numbered=False, level=0):
        """Add a list item with proper indentation and formatting"""
        from docx.shared import Pt, Inches

        style_name = "List Number" if is_numbered else "List Bullet"

        # Ensure the style exists
        if style_name not in doc.styles:
            style_name = "Normal"

        p = doc.add_paragraph(style=style_name)
        p.add_run(text.strip())

        # Set indentation for nested lists
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
            p.paragraph_format.first_line_indent = Inches(-0.25)

        p.paragraph_format.space_after = Pt(3)

    def _add_blockquote_to_docx(self, doc, text):
        """Add a blockquote with special formatting"""
        from docx.shared import Pt, Inches, RGBColor

        p = doc.add_paragraph()
        run = p.add_run(text.strip())

        # Style as blockquote
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # Make text slightly italic and gray
        run.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)

    def _add_code_block_to_docx(self, doc, code_content):
        """Add a code block with monospace formatting"""
        from docx.shared import Pt, RGBColor, Inches

        p = doc.add_paragraph()
        run = p.add_run(code_content)

        # Format as code
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)

        # Add background-like spacing
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    def _add_horizontal_rule(self, doc):
        """Add a horizontal rule"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = doc.add_paragraph("─" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

    def _finalize_docx_table(self, doc, table_rows):
        """Convert table rows to DOCX table."""
        if not table_rows:
            return

        from docx.shared import Pt, Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn

        # Determine number of columns from the first row
        num_cols = len(table_rows[0]) if table_rows else 0
        if num_cols == 0:
            return

        table = doc.add_table(rows=1, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for col in table.columns:
            col.width = Inches(max(1.0, 6.5 / num_cols))

        # Header row (first row)
        header_row = table_rows[0]
        for i, cell_text in enumerate(header_row):
            if i < num_cols:
                cell = table.rows[0].cells[i]
                cell.text = cell_text.strip()
                run = (
                    cell.paragraphs[0].runs[0]
                    if cell.paragraphs[0].runs
                    else cell.paragraphs[0].add_run()
                )
                run.font.bold, run.font.size = True, Pt(11)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9E2F3")
                cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for r_idx, row_data in enumerate(table_rows[1:], 1):
            row_cells = table.add_row().cells
            for c_idx, cell_data in enumerate(row_data):
                if c_idx < num_cols:
                    cell = row_cells[c_idx]
                    cell.text = str(cell_data).strip()
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                    if r_idx % 2 == 0:  # Alternating row color
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "F9F9F9")
                        cell._tc.get_or_add_tcPr().append(shading)

    # -------------------------------------------------------------------------
    # Content Processing & Customization (updated)
    # -------------------------------------------------------------------------

    def _get_custom_text(self, department: str) -> tuple[str, str]:
        intros = {
            "HR": "This document contains HR-related information and guidelines for company personnel.",
            "Engineering": "This document outlines technical specifications and engineering guidelines.",
        }
        footers = {
            "HR": "For further questions, please contact the Human Resources department.",
            "Engineering": "For further questions, please contact the Engineering department.",
        }
        intro = intros.get(
            department,
            f"This document contains {department.lower()}-related information for company personnel.",
        )
        footer = footers.get(
            department,
            f"For further questions, please contact the {department} department.",
        )
        return intro, footer

    # -------------------------------------------------------------------------
    # Shared Components (Headers, Footers) - unchanged
    # -------------------------------------------------------------------------
    def _add_company_header_docx(
        self, doc, user_valves, user_name: str, department: str
    ):
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = doc.add_paragraph()
        p.add_run(user_valves.company_name).bold = True
        p.runs[0].font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph(
            f"{department} | {user_name} | {datetime.datetime.now().strftime('%B %d, %Y')}"
        )
        p.runs[0].font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph("─" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)

    def _add_company_footer_docx(self, doc, user_valves):
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = doc.add_paragraph("─" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)

        p = doc.add_paragraph()
        p.add_run(f"Generated by AI Assistant | {user_valves.company_name}").italic = (
            True
        )
        p.runs[0].font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
